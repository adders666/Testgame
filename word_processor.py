import tkinter as tk
from tkinter import filedialog, messagebox
import os

from docx import Document
import PyPDF2
from bs4 import BeautifulSoup

class WordProcessor:
    def __init__(self, root):
        self.root = root
        root.title("Word Processor")
        self.filename = None

        self.text = tk.Text(root, wrap='word')
        self.text.pack(fill='both', expand=True)

        self.menu = tk.Menu(root)
        root.config(menu=self.menu)

        file_menu = tk.Menu(self.menu, tearoff=0)
        file_menu.add_command(label="New", command=self.new_file)
        file_menu.add_command(label="Open", command=self.open_file)
        file_menu.add_command(label="Save", command=self.save_file)
        file_menu.add_command(label="Save As", command=self.save_as_file)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=root.quit)
        self.menu.add_cascade(label="File", menu=file_menu)

        insert_menu = tk.Menu(self.menu, tearoff=0)
        insert_menu.add_command(label="Page Break", command=self.insert_page_break)
        insert_menu.add_command(label="Insert PDF", command=self.insert_pdf)
        insert_menu.add_command(label="Insert HTML", command=self.insert_html)
        self.menu.add_cascade(label="Insert", menu=insert_menu)

        self.toolbar = tk.Frame(root)
        self.bold_btn = tk.Button(self.toolbar, text="Bold", command=self.make_bold)
        self.bold_btn.pack(side='left')
        self.italic_btn = tk.Button(self.toolbar, text="Italic", command=self.make_italic)
        self.italic_btn.pack(side='left')
        self.toolbar.pack(fill='x')

        default_font = ('Arial', 12)
        self.text.configure(font=default_font)
        self.text.tag_configure("bold", font=('Arial', 12, 'bold'))
        self.text.tag_configure("italic", font=('Arial', 12, 'italic'))

    def make_bold(self):
        self.apply_tag("bold")

    def make_italic(self):
        self.apply_tag("italic")

    def apply_tag(self, tag):
        try:
            start = self.text.index("sel.first")
            end = self.text.index("sel.last")
            self.text.tag_add(tag, start, end)
        except tk.TclError:
            pass

    def new_file(self):
        self.text.delete(1.0, tk.END)
        self.filename = None

    def open_file(self):
        filetypes = [
            ("All files", "*.*"),
            ("Text files", "*.txt"),
            ("Rich Text Format", "*.rtf"),
            ("Word Documents", "*.docx"),
            ("PDF", "*.pdf"),
            ("HTML", "*.html")
        ]
        fname = filedialog.askopenfilename(filetypes=filetypes)
        if not fname:
            return
        self.filename = fname
        ext = os.path.splitext(fname)[1].lower()
        if ext == '.txt':
            with open(fname, 'r') as f:
                data = f.read()
        elif ext == '.rtf':
            from striprtf.striprtf import rtf_to_text
            with open(fname, 'r') as f:
                data = rtf_to_text(f.read())
        elif ext == '.docx':
            doc = Document(fname)
            data = '\n'.join(p.text for p in doc.paragraphs)
        elif ext == '.pdf':
            with open(fname, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                data = ''
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        data += text + '\n'
        elif ext == '.html':
            with open(fname, 'r') as f:
                soup = BeautifulSoup(f, 'html.parser')
                data = soup.get_text()
        else:
            messagebox.showerror("Unsupported format", "Cannot open this file type")
            return
        self.text.delete(1.0, tk.END)
        self.text.insert(tk.END, data)

    def save_file(self):
        if self.filename:
            self.save_to_file(self.filename)
        else:
            self.save_as_file()

    def save_as_file(self):
        filetypes = [
            ("Text files", "*.txt"),
            ("Rich Text Format", "*.rtf"),
            ("Word Documents", "*.docx"),
            ("PDF", "*.pdf")
        ]
        fname = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=filetypes)
        if fname:
            self.filename = fname
            self.save_to_file(fname)

    def save_to_file(self, fname):
        ext = os.path.splitext(fname)[1].lower()
        text = self.text.get(1.0, tk.END)
        if ext == '.txt':
            with open(fname, 'w') as f:
                f.write(text)
        elif ext == '.rtf':
            from striprtf.striprtf import text_to_rtf
            rtf = text_to_rtf(text)
            with open(fname, 'w') as f:
                f.write(rtf)
        elif ext == '.docx':
            doc = Document()
            for line in text.splitlines():
                doc.add_paragraph(line)
            doc.save(fname)
        elif ext == '.pdf':
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            c = canvas.Canvas(fname, pagesize=letter)
            width, height = letter
            y = height - 50
            for line in text.splitlines():
                c.drawString(50, y, line)
                y -= 15
                if y < 50:
                    c.showPage()
                    y = height - 50
            c.save()
        else:
            messagebox.showerror("Unsupported format", "Cannot save this file type")

    def insert_page_break(self):
        self.text.insert(tk.INSERT, "\n------- PAGE BREAK -------\n")

    def insert_pdf(self):
        fname = filedialog.askopenfilename(filetypes=[("PDF", "*.pdf")])
        if not fname:
            return
        with open(fname, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            data = ''
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    data += text + '\n'
        self.text.insert(tk.INSERT, data)

    def insert_html(self):
        fname = filedialog.askopenfilename(filetypes=[("HTML files", "*.html"), ("HTM files", "*.htm")])
        if not fname:
            return
        with open(fname, 'r') as f:
            soup = BeautifulSoup(f, 'html.parser')
            data = soup.get_text()
        self.text.insert(tk.INSERT, data)

if __name__ == '__main__':
    root = tk.Tk()
    app = WordProcessor(root)
    root.mainloop()
